import cv2
import pytesseract
from PIL import Image
import docx
import io
import logging
from typing import Dict, Optional
from pathlib import Path
import fitz  # PyMuPDF for better PDF handling

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class DocumentProcessor:
    """Enhanced document processor supporting multiple file formats."""
    
    SUPPORTED_FORMATS = {
        'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp'],
        'pdf': ['.pdf'],
        'text': ['.txt'],
        'word': ['.docx', '.doc']
    }

    def __init__(self, ocr_config: Optional[Dict] = None):
        self.logger = logging.getLogger(__name__)
        self.ocr_config = ocr_config or {}
        
        # Configure Tesseract
        if 'tesseract_cmd' in self.ocr_config:
            pytesseract.pytesseract.tesseract_cmd = self.ocr_config['tesseract_cmd']
            
        # Initialize OCR languages
        self.ocr_langs = self.ocr_config.get('languages', 'eng')

    def extract_text(self, file_path: str) -> Dict[str, str]:
        """
        Extract text from various file formats.
        Returns a dictionary with metadata and extracted text.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()
        
        try:
            # Determine file type and call appropriate extractor
            if extension in self.SUPPORTED_FORMATS['image']:
                return self._extract_from_image(file_path)
            elif extension in self.SUPPORTED_FORMATS['pdf']:
                return self._extract_from_pdf(file_path)
            elif extension in self.SUPPORTED_FORMATS['text']:
                return self._extract_from_text(file_path)
            elif extension in self.SUPPORTED_FORMATS['word']:
                return self._extract_from_word(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
                
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            raise

    def _extract_from_image(self, file_path: Path) -> Dict[str, str]:
        """Extract text from image files using OCR."""
        try:
            # Open image with PIL for better format support
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode not in ('L', 'RGB'):
                image = image.convert('RGB')
            
            # Perform OCR with configured languages
            text = pytesseract.image_to_string(
                image, 
                lang=self.ocr_langs,
                config=self.ocr_config.get('tesseract_config', '')
            )
            
            # Get image metadata
            metadata = {
                'format': image.format,
                'size': image.size,
                'mode': image.mode
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'type': 'image'
            }
            
        except Exception as e:
            self.logger.error(f"Error in OCR processing: {str(e)}")
            raise

    def _extract_from_pdf(self, file_path: Path) -> Dict[str, str]:
        """Extract text from PDF files with enhanced handling."""
        try:
            # Use PyMuPDF (fitz) for better PDF handling
            doc = fitz.open(file_path)
            
            pages_text = []
            images = []
            metadata = {
                'page_count': len(doc),
                'author': doc.metadata.get('author', ''),
                'title': doc.metadata.get('title', ''),
                'subject': doc.metadata.get('subject', ''),
                'keywords': doc.metadata.get('keywords', '')
            }

            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text
                pages_text.append(page.get_text())
                
                # Extract images if present
                image_list = page.get_images()
                if image_list:
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # Use PIL to process the image
                            image = Image.open(io.BytesIO(image_bytes))
                            ocr_text = pytesseract.image_to_string(
                                image,
                                lang=self.ocr_langs
                            )
                            pages_text.append(ocr_text)
                            
                        except Exception as img_error:
                            self.logger.warning(f"Error processing image {img_index} on page {page_num}: {str(img_error)}")

            return {
                'text': '\n'.join(pages_text),
                'metadata': metadata,
                'type': 'pdf'
            }

        except Exception as e:
            self.logger.error(f"Error in PDF processing: {str(e)}")
            raise

    def _extract_from_text(self, file_path: Path) -> Dict[str, str]:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'ascii']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
                    
            if text is None:
                raise ValueError(f"Could not decode file with encodings: {encodings}")
                
            return {
                'text': text,
                'metadata': {
                    'encoding': encoding,
                    'size': file_path.stat().st_size
                },
                'type': 'text'
            }
            
        except Exception as e:
            self.logger.error(f"Error in text file processing: {str(e)}")
            raise

    def _extract_from_word(self, file_path: Path) -> Dict[str, str]:
        """Extract text from Word documents."""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(' | '.join(row_text))
            
            # Get document properties
            metadata = {
                'core_properties': {
                    'author': doc.core_properties.author,
                    'created': str(doc.core_properties.created),
                    'modified': str(doc.core_properties.modified),
                    'title': doc.core_properties.title
                }
            }
            
            return {
                'text': '\n'.join(full_text),
                'metadata': metadata,
                'type': 'word'
            }
            
        except Exception as e:
            self.logger.error(f"Error in Word document processing: {str(e)}")
            raise

    def preprocess_image(self, image):
        """Apply image preprocessing for better OCR results."""
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply thresholding
        _, threshold = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Noise removal
        denoised = cv2.fastNlMeansDenoising(threshold)
        
        return denoised

# Usage example
if __name__ == "__main__":
    # Configure OCR settings
    ocr_config = {
        'languages': 'eng+fra+deu',  # Multiple language support
        'tesseract_config': '--oem 3 --psm 6'
    }
    
    processor = DocumentProcessor(ocr_config=ocr_config)
    
    # Test with different file types
    test_files = [
        'sample.pdf',
        'document.docx',
        'image.jpg',
        'text.txt'
    ]
    
    for file_path in test_files:
        try:
            result = processor.extract_text(file_path)
            print(f"\nProcessing {file_path}:")
            print(f"Type: {result['type']}")
            print(f"Metadata: {result['metadata']}")
            print(f"First 100 chars: {result['text'][:100]}...")
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")